import logging
from io import BytesIO

import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

def extract_text_from_pdf(data: bytes) -> str:
    """
    Extracts text from a PDF byte stream.
    Validates PDF header and handles corrupted files gracefully.
    """
    # 1. Validate PDF Header (%PDF-)
    if not data.startswith(b"%PDF-"):
        logger.warning("File rejection: Missing %PDF- header")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file format: Not a valid PDF (missing %PDF- header)."
        )

    try:
        with pdfplumber.open(BytesIO(data)) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            extracted_text = "\n\n".join(text_parts).strip()
            
            if not extracted_text:
                logger.info("PDF extraction: No text found (might be an image-only PDF)")
                # Fallback to pypdf or just return empty if OCR is not available
                return ""
                
            return extracted_text

    except Exception as e:
        logger.error(f"Failed to parse PDF: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse PDF: The file might be corrupted or encrypted. Error: {str(e)}"
        )

def extract_text_from_docx(data: bytes) -> str:
    """Extract plain text from a Word .docx (Office Open XML) file."""
    if len(data) < 4 or data[:2] != b"PK":
        logger.warning("File rejection: DOCX should be a ZIP-based OOXML document")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid file format: Not a valid DOCX (expected ZIP package signature).",
        )
    try:
        doc = Document(BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts).strip()
        if not text:
            logger.info("DOCX extraction: no text in paragraphs or tables")
        return text
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Failed to parse DOCX: %s", e, exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to parse DOCX: {e!s}",
        ) from e


def extract_text_from_html(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return " ".join(soup.get_text(separator=" ").split())
